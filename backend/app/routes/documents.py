import io
import docx
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException
from sqlalchemy.orm import Session
from sqlalchemy import text
from app.database import get_db
from app.models import Document, DocumentTopic
from collections import defaultdict
from langchain_text_splitters import RecursiveCharacterTextSplitter

router = APIRouter()


@router.get("/documents")
def list_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).order_by(Document.uploaded_at.desc()).all()
    return {
        "documents": [
            {
                "id": d.id,
                "filename": d.filename,
                "subject": d.subject,
                "uploaded_at": d.uploaded_at.isoformat() if d.uploaded_at else None,
            }
            for d in docs
        ]
    }


@router.get("/documents/hierarchy")
def get_document_hierarchy(db: Session = Depends(get_db)):
    rows = db.query(DocumentTopic).all()
    hierarchy = defaultdict(lambda: defaultdict(lambda: defaultdict(set)))

    for row in rows:
        subject = row.subject or "General"
        doc_label = f"{row.filename} (ID: {row.document_id})"
        topic = row.topic or row.section or "General"
        subtopic = row.subtopic or "General"
        hierarchy[subject][doc_label][topic].add(subtopic)

    result = []
    for subject, docs in hierarchy.items():
        doc_items = []
        for doc_label, topics in docs.items():
            topic_items = []
            for topic, subtopics in topics.items():
                topic_items.append({
                    "topic": topic,
                    "subtopics": sorted(subtopics)
                })
            doc_items.append({
                "document": doc_label,
                "topics": sorted(topic_items, key=lambda t: t["topic"])
            })
        result.append({
            "subject": subject,
            "documents": sorted(doc_items, key=lambda d: d["document"])
        })

    return {"hierarchy": sorted(result, key=lambda r: r["subject"])}


@router.post("/api/documents/upload")
async def upload_document_api(
    file: UploadFile = File(...),
    subject: str = Form("General"),
    db: Session = Depends(get_db)
):
    filename = file.filename
    if not (filename.lower().endswith(".pdf") or filename.lower().endswith(".docx") or filename.lower().endswith(".txt")):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, DOCX, and TXT are supported.")

    try:
        file_content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file content: {str(e)}")

    # Ensure PostgreSQL table exists
    try:
        db.execute(text("""
            CREATE TABLE IF NOT EXISTS document_chunks (
                id SERIAL PRIMARY KEY,
                doc_id INTEGER,
                content TEXT,
                page_num INTEGER,
                chunk_index INTEGER,
                filename VARCHAR(255)
            );
        """))
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database table initialization failed: {str(e)}")

    # Insert document record in Postgres
    try:
        db_doc = Document(filename=filename, subject=subject)
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
        doc_id = db_doc.id
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to save document record: {str(e)}")

    chunks_to_insert = []
    
    # Text splitter config
    try:
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            chunk_size=512,
            chunk_overlap=50
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to initialize text splitter: {str(e)}")

    # Extract text and chunk
    try:
        if filename.lower().endswith(".pdf"):
            pdf_file = io.BytesIO(file_content)
            extracted_pages = []
            
            # 1. Try pypdf
            try:
                import pypdf
                reader = pypdf.PdfReader(pdf_file)
                for page_idx, page in enumerate(reader.pages):
                    text_content = page.extract_text() or ""
                    extracted_pages.append((page_idx + 1, text_content))
                print(f"[documents] Successfully extracted text using pypdf: {len(extracted_pages)} pages")
            except Exception as e_pypdf:
                print(f"[documents] pypdf failed: {e_pypdf}")
                extracted_pages = []

            # 2. Try fitz (PyMuPDF) if pypdf failed or got no pages
            if not extracted_pages:
                try:
                    import fitz
                    doc = fitz.open(stream=file_content, filetype="pdf")
                    for page_idx, page in enumerate(doc):
                        text_content = page.get_text() or ""
                        extracted_pages.append((page_idx + 1, text_content))
                    print(f"[documents] Successfully extracted text using fitz: {len(extracted_pages)} pages")
                except Exception as e_fitz:
                    print(f"[documents] fitz failed: {e_fitz}")
                    extracted_pages = []

            # 3. Retry pypdf with seek reset if both fitz and first pypdf failed
            if not extracted_pages:
                try:
                    import pypdf as _pypdf2_retry
                    pdf_file.seek(0)
                    reader = _pypdf2_retry.PdfReader(pdf_file)
                    for page_idx, page in enumerate(reader.pages):
                        text_content = page.extract_text() or ""
                        extracted_pages.append((page_idx + 1, text_content))
                    print(f"[documents] Successfully extracted text (pypdf retry): {len(extracted_pages)} pages")
                except Exception as e_retry:
                    print(f"[documents] pypdf retry failed: {e_retry}")
                    extracted_pages = []

            # 4. If extracted text is empty/meaningless (scanned/image-only PDF), fallback to Gemini OCR
            total_text_len = sum(len(text) for _, text in extracted_pages)
            if total_text_len < 10:
                print(f"[documents] Local extraction got only {total_text_len} characters. Performing Gemini Multimodal OCR...")
                try:
                    import os
                    import re
                    import fitz
                    from google import genai
                    from google.genai import types
                    
                    gemini_api_key = os.getenv("GEMINI_API_KEY")
                    if gemini_api_key:
                        client = genai.Client(api_key=gemini_api_key)
                        
                        # Call Gemini to OCR the entire PDF in a single call (fast, cost-effective, rate-limit friendly)
                        response = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=[
                                types.Part.from_bytes(
                                    data=file_content,
                                    mime_type='application/pdf',
                                ),
                                'Perform OCR on this scanned PDF and extract all text content page by page. '
                                'Clearly separate pages with markers like "--- PAGE 1 ---", "--- PAGE 2 ---", etc. '
                                'Extract everything verbatim.',
                            ]
                        )
                        full_text = response.text or ""
                        
                        # Parse page by page using regex split
                        pattern = re.compile(
                            r'(?:---|###|#)?\s*(?:PAGE|Page)\s*[:#-]?\s*(\d+)\s*(?:of\s*\d+)?\s*(?:---|###|#)?',
                            re.IGNORECASE
                        )
                        matches = list(pattern.finditer(full_text))
                        ocr_pages = []
                        
                        if not matches:
                            # If no page markers are found, distribute the text evenly
                            doc = fitz.open(stream=file_content, filetype="pdf")
                            total_pages = len(doc)
                            chunk_len = len(full_text) // total_pages if total_pages > 0 else len(full_text)
                            for i in range(total_pages):
                                start = i * chunk_len
                                end = start + chunk_len if i < total_pages - 1 else len(full_text)
                                ocr_pages.append((i + 1, full_text[start:end]))
                        else:
                            for idx, match in enumerate(matches):
                                try:
                                    page_num = int(match.group(1))
                                except ValueError:
                                    page_num = idx + 1
                                    
                                start_idx = match.end()
                                if idx + 1 < len(matches):
                                    end_idx = matches[idx + 1].start()
                                else:
                                    end_idx = len(full_text)
                                    
                                page_text = full_text[start_idx:end_idx].strip()
                                ocr_pages.append((page_num, page_text))
                        
                        if ocr_pages:
                            ocr_pages.sort(key=lambda x: x[0])
                            extracted_pages = ocr_pages
                            print(f"[documents] Gemini OCR successfully extracted text for {len(extracted_pages)} pages!")
                except Exception as ocr_global_e:
                    print(f"[documents] Gemini OCR setup failed: {ocr_global_e}")

            chunk_idx = 0
            for page_num, page_text in extracted_pages:
                if not page_text.strip():
                    continue
                page_chunks = splitter.split_text(page_text)
                for chunk in page_chunks:
                    chunks_to_insert.append({
                        "doc_id": doc_id,
                        "content": chunk,
                        "page_num": page_num,
                        "chunk_index": chunk_idx,
                        "filename": filename
                    })
                    chunk_idx += 1

        elif filename.lower().endswith(".docx"):
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            doc_chunks = splitter.split_text(full_text)
            for chunk_idx, chunk in enumerate(doc_chunks):
                chunks_to_insert.append({
                    "doc_id": doc_id,
                    "content": chunk,
                    "page_num": 1,
                    "chunk_index": chunk_idx,
                    "filename": filename
                })

        else:  # .txt
            full_text = file_content.decode("utf-8", errors="ignore")
            doc_chunks = splitter.split_text(full_text)
            for chunk_idx, chunk in enumerate(doc_chunks):
                chunks_to_insert.append({
                    "doc_id": doc_id,
                    "content": chunk,
                    "page_num": 1,
                    "chunk_index": chunk_idx,
                    "filename": filename
                })
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Text extraction failed: {str(e)}")

    # Store chunks in PostgreSQL
    if chunks_to_insert:
        try:
            for chunk in chunks_to_insert:
                db.execute(text("""
                    INSERT INTO document_chunks (doc_id, content, page_num, chunk_index, filename)
                    VALUES (:doc_id, :content, :page_num, :chunk_index, :filename)
                """), chunk)
            db.commit()
        except Exception as e:
            import traceback
            traceback.print_exc()
            db.rollback()
            raise HTTPException(status_code=500, detail=f"Failed to store chunks: {str(e)}")

        # Store in Pinecone vector DB (non-fatal — upload still succeeds if Pinecone fails)
        pinecone_ok = False
        try:
            try:
                from app.services.embeddings import embed_and_store
            except ImportError:
                try:
                    from services.embeddings import embed_and_store
                except ImportError:
                    from backend.services.embeddings import embed_and_store
            
            chunks_text = [c["content"] for c in chunks_to_insert]
            embed_and_store(chunks_text, str(doc_id), filename)
            pinecone_ok = True
            print(f"[UPLOAD] Pinecone vector storage successful for doc_id={doc_id}")
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[UPLOAD] WARNING: Pinecone vector storage failed (non-fatal): {e}")
            print(f"[UPLOAD] Document saved to PostgreSQL — chat will use DB full-text search fallback.")

    return {
        "doc_id": doc_id,
        "chunk_count": len(chunks_to_insert),
        "vector_indexed": pinecone_ok if chunks_to_insert else False
    }


@router.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    """Delete a document and all its chunks from PostgreSQL and Pinecone."""
    # Check document exists
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")

    # Delete from Pinecone
    try:
        try:
            from app.services.embeddings import get_pinecone_index
        except ImportError:
            try:
                from services.embeddings import get_pinecone_index
            except ImportError:
                from backend.services.embeddings import get_pinecone_index

        index = get_pinecone_index()
        # Delete all vectors with matching doc_id metadata filter
        index.delete(filter={"doc_id": {"$eq": str(doc_id)}})
        print(f"[DELETE] Deleted Pinecone vectors for doc_id={doc_id}")
    except Exception as e:
        print(f"[DELETE] Pinecone deletion warning (non-fatal): {e}")

    # Delete from document_chunks
    try:
        db.execute(text("DELETE FROM document_chunks WHERE doc_id = :doc_id"), {"doc_id": doc_id})
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete chunks: {str(e)}")

    # Delete from documents table
    try:
        db.delete(doc)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

    return {"success": True, "deleted_doc_id": doc_id}

